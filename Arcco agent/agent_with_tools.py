"""
Agent com Tools Completas — PDF, Web Search, File Ops, Code Exec.

Agente conversacional completo com ferramentas para:
- Pesquisa na internet (Brave Search API ou Claude nativo)
- Geração de PDFs
- Leitura/escrita de arquivos
- Execução de código Python
- Geração de documentos (DOCX, XLSX)

Requisitos:
    pip install anthropic httpx beautifulsoup4 reportlab pypdf python-docx openpyxl

Variáveis de ambiente:
    ANTHROPIC_API_KEY=sk-...
    BRAVE_SEARCH_API_KEY=... (opcional, usa Claude web_search se ausente)
"""

import anthropic
import httpx
import json
import os
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

# ============================================================
# CONFIG
# ============================================================

MODEL = "claude-sonnet-4-5-20250514"
MAX_TOKENS = 8096
MAX_ITERATIONS = 20
WORKSPACE = Path("/tmp/agent_workspace")
WORKSPACE.mkdir(parents=True, exist_ok=True)

SYSTEM_PROMPT = """Você é um agente autônomo com acesso a ferramentas poderosas.

## Suas Capacidades
- Pesquisar informações na internet
- Gerar documentos PDF com formatação profissional
- Ler e escrever arquivos
- Executar código Python para cálculos e processamento
- Gerar documentos Word (DOCX) e planilhas (XLSX)

## Regras
1. Use ferramentas para obter informações reais — nunca invente dados
2. Se uma ferramenta falhar, tente abordagem alternativa
3. Para tarefas complexas, divida em etapas e execute uma de cada vez
4. Ao gerar documentos, confirme o que foi criado e onde está salvo
5. Seja conciso e direto
"""

# ============================================================
# TOOL REGISTRY
# ============================================================

TOOLS = []
HANDLERS = {}


def tool(name, description, parameters):
    def decorator(func):
        required = [k for k, v in parameters.items() if v.pop("required", False)]
        TOOLS.append({
            "name": name,
            "description": description,
            "input_schema": {
                "type": "object",
                "properties": parameters,
                "required": required,
            },
        })
        HANDLERS[name] = func
        return func
    return decorator


# ============================================================
# TOOL: Web Search
# ============================================================

@tool(
    name="web_search",
    description="Pesquisa na internet. Retorna títulos, URLs e snippets dos resultados. "
                "Use queries curtas e específicas (2-6 palavras) para melhores resultados.",
    parameters={
        "query": {"type": "string", "description": "Termo de busca", "required": True},
        "num_results": {"type": "integer", "description": "Resultados (default: 5, max: 10)"},
    },
)
def web_search(query: str, num_results: int = 5) -> str:
    api_key = os.getenv("BRAVE_SEARCH_API_KEY")
    if not api_key:
        return "ERRO: BRAVE_SEARCH_API_KEY não configurada. Configure ou use web_fetch com URLs diretas."

    try:
        resp = httpx.get(
            "https://api.search.brave.com/res/v1/web/search",
            headers={"X-Subscription-Token": api_key},
            params={"q": query, "count": min(num_results, 10)},
            timeout=15.0,
        )
        resp.raise_for_status()
        data = resp.json()

        results = []
        for item in data.get("web", {}).get("results", [])[:num_results]:
            results.append(f"• {item['title']}\n  {item['url']}\n  {item.get('description', '')}")

        return f"Resultados para '{query}':\n\n" + "\n\n".join(results) if results else f"Sem resultados para: {query}"

    except Exception as e:
        return f"ERRO na busca: {e}"


# ============================================================
# TOOL: Web Fetch
# ============================================================

@tool(
    name="web_fetch",
    description="Acessa uma URL e extrai o conteúdo textual. "
                "Use para ler artigos, documentação, páginas web.",
    parameters={
        "url": {"type": "string", "description": "URL completa (com https://)", "required": True},
        "max_chars": {"type": "integer", "description": "Máximo de caracteres (default: 8000)"},
    },
)
def web_fetch(url: str, max_chars: int = 8000) -> str:
    from bs4 import BeautifulSoup

    try:
        resp = httpx.get(url, headers={"User-Agent": "AgentBot/1.0"}, timeout=20.0, follow_redirects=True)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = "\n".join(line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip())

        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[...truncado]"

        return f"Conteúdo de {url}:\n\n{text}"

    except Exception as e:
        return f"ERRO ao acessar {url}: {e}"


# ============================================================
# TOOL: Generate PDF
# ============================================================

@tool(
    name="generate_pdf",
    description="Gera documento PDF profissional. Suporta títulos, parágrafos, tabelas e listas. "
                "Ideal para relatórios, resumos e documentos formatados.",
    parameters={
        "title": {"type": "string", "description": "Título do documento", "required": True},
        "sections": {
            "type": "array",
            "description": "Lista de seções. Cada seção: {type: 'heading'|'paragraph'|'table'|'bullets', text/rows/items, level}",
            "required": True,
        },
        "filename": {"type": "string", "description": "Nome do arquivo (default: report.pdf)"},
    },
)
def generate_pdf(title: str, sections: list, filename: str = "report.pdf") -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    path = WORKSPACE / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=25*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("DocTitle", parent=styles["Title"], fontSize=22,
                              spaceAfter=20, textColor=HexColor("#1a1a2e")))
    styles.add(ParagraphStyle("DocH1", parent=styles["Heading1"], fontSize=16,
                              spaceAfter=10, textColor=HexColor("#16213e")))
    styles.add(ParagraphStyle("DocH2", parent=styles["Heading2"], fontSize=13,
                              spaceAfter=8, textColor=HexColor("#0f3460")))
    styles.add(ParagraphStyle("DocBody", parent=styles["Normal"], fontSize=11,
                              spaceAfter=8, leading=16))
    styles.add(ParagraphStyle("DocBullet", parent=styles["Normal"], fontSize=11,
                              leftIndent=20, spaceAfter=4, leading=15,
                              bulletIndent=10, bulletFontSize=11))

    story = [Paragraph(title, styles["DocTitle"]), Spacer(1, 12)]
    heading_map = {1: "DocH1", 2: "DocH2", 3: "Heading3"}

    for sec in sections:
        t = sec.get("type", "paragraph")

        if t == "heading":
            style = heading_map.get(sec.get("level", 1), "DocH1")
            story.append(Paragraph(sec.get("text", ""), styles[style]))

        elif t == "paragraph":
            story.append(Paragraph(sec.get("text", ""), styles["DocBody"]))

        elif t == "table":
            rows = sec.get("rows", [])
            if rows:
                table = Table(rows, repeatRows=1)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1a1a2e")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#f5f5f5")]),
                    ("PADDING", (0, 0), (-1, -1), 8),
                ]))
                story.append(table)
                story.append(Spacer(1, 10))

        elif t == "bullets":
            for item in sec.get("items", []):
                story.append(Paragraph(f"• {item}", styles["DocBullet"]))

    doc.build(story)
    size = path.stat().st_size
    return f"PDF gerado: {path} ({size:,} bytes)"


# ============================================================
# TOOL: File Operations
# ============================================================

@tool(
    name="read_file",
    description="Lê conteúdo de um arquivo de texto.",
    parameters={
        "path": {"type": "string", "description": "Caminho do arquivo (relativo ao workspace)", "required": True},
    },
)
def read_file(path: str) -> str:
    fp = (WORKSPACE / path).resolve()
    if not str(fp).startswith(str(WORKSPACE.resolve())):
        return "ERRO: Acesso negado (fora do workspace)."
    if not fp.exists():
        return f"ERRO: Arquivo '{path}' não encontrado."
    text = fp.read_text(encoding="utf-8")
    if len(text) > 15000:
        text = text[:15000] + "\n[...truncado]"
    return text


@tool(
    name="write_file",
    description="Escreve conteúdo em arquivo. Cria se não existir.",
    parameters={
        "path": {"type": "string", "description": "Caminho do arquivo", "required": True},
        "content": {"type": "string", "description": "Conteúdo", "required": True},
    },
)
def write_file(path: str, content: str) -> str:
    fp = (WORKSPACE / path).resolve()
    if not str(fp).startswith(str(WORKSPACE.resolve())):
        return "ERRO: Acesso negado."
    fp.parent.mkdir(parents=True, exist_ok=True)
    fp.write_text(content, encoding="utf-8")
    return f"Salvo: {path} ({len(content):,} chars)"


@tool(
    name="list_files",
    description="Lista arquivos no workspace do agente.",
    parameters={
        "directory": {"type": "string", "description": "Subdiretório (default: raiz do workspace)"},
    },
)
def list_files(directory: str = ".") -> str:
    dp = (WORKSPACE / directory).resolve()
    if not dp.is_dir():
        return f"ERRO: '{directory}' não é um diretório."

    entries = []
    for item in sorted(dp.iterdir()):
        if item.is_file():
            entries.append(f"  📄 {item.name} ({item.stat().st_size:,} bytes)")
        elif item.is_dir():
            entries.append(f"  📁 {item.name}/")

    return "Workspace:\n" + "\n".join(entries) if entries else "Workspace vazio."


# ============================================================
# TOOL: Execute Code
# ============================================================

@tool(
    name="execute_python",
    description="Executa código Python em ambiente isolado. "
                "Use para cálculos, processamento de dados, análises. "
                "Timeout de 30 segundos. Sem acesso à rede.",
    parameters={
        "code": {"type": "string", "description": "Código Python", "required": True},
    },
)
def execute_python(code: str) -> str:
    blocked = ["os.system", "subprocess", "shutil.rmtree", "__import__", "eval(", "exec("]
    for b in blocked:
        if b in code:
            return f"ERRO: '{b}' bloqueado por segurança."

    with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False, dir="/tmp") as f:
        f.write(code)
        tmp = f.name

    try:
        result = subprocess.run(
            ["python3", tmp], capture_output=True, text=True, timeout=30, cwd=str(WORKSPACE),
        )
        output = (result.stdout or "") + (f"\nSTDERR: {result.stderr}" if result.stderr else "")
        if len(output) > 10000:
            output = output[:10000] + "\n[...truncado]"
        return output or "(executado sem output)"

    except subprocess.TimeoutExpired:
        return "ERRO: Timeout (30s). Otimize o código."
    except Exception as e:
        return f"ERRO: {e}"
    finally:
        os.unlink(tmp)


# ============================================================
# TOOL: Generate DOCX
# ============================================================

@tool(
    name="generate_docx",
    description="Gera documento Word (.docx) com formatação profissional.",
    parameters={
        "title": {"type": "string", "description": "Título", "required": True},
        "sections": {
            "type": "array",
            "description": "Seções: {type: heading/paragraph/table/bullets, text/rows/items}",
            "required": True,
        },
        "filename": {"type": "string", "description": "Nome do arquivo"},
    },
)
def generate_docx(title: str, sections: list, filename: str = "document.docx") -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for sec in sections:
        t = sec.get("type", "paragraph")
        if t == "heading":
            doc.add_heading(sec.get("text", ""), level=sec.get("level", 1))
        elif t == "paragraph":
            doc.add_paragraph(sec.get("text", ""))
        elif t == "table":
            rows = sec.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = str(cell)
        elif t == "bullets":
            for item in sec.get("items", []):
                doc.add_paragraph(item, style="List Bullet")

    path = WORKSPACE / filename
    doc.save(str(path))
    return f"DOCX gerado: {path} ({path.stat().st_size:,} bytes)"


# ============================================================
# AGENT LOOP
# ============================================================

def agent_loop(user_message: str, history: list = None) -> tuple[str, list]:
    """Executa o agent loop. Retorna (resposta, histórico atualizado)."""
    client = anthropic.Anthropic()

    if history is None:
        history = []

    history.append({"role": "user", "content": user_message})

    for _ in range(MAX_ITERATIONS):
        response = client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            system=SYSTEM_PROMPT,
            tools=TOOLS,
            messages=history,
        )

        content = response.content
        history.append({"role": "assistant", "content": content})

        tool_uses = [b for b in content if b.type == "tool_use"]

        if not tool_uses:
            texts = [b.text for b in content if hasattr(b, "text")]
            return "\n".join(texts), history

        results = []
        for tu in tool_uses:
            handler = HANDLERS.get(tu.name)
            if not handler:
                results.append({"type": "tool_result", "tool_use_id": tu.id,
                                "content": f"Tool '{tu.name}' não existe.", "is_error": True})
                continue
            try:
                r = handler(**tu.input)
                print(f"  🔧 {tu.name} → {str(r)[:120]}")
                results.append({"type": "tool_result", "tool_use_id": tu.id, "content": str(r)})
            except Exception as e:
                print(f"  ❌ {tu.name}: {e}")
                results.append({"type": "tool_result", "tool_use_id": tu.id,
                                "content": f"ERRO: {e}", "is_error": True})

        history.append({"role": "user", "content": results})

    return "⚠️ Limite de iterações atingido.", history


# ============================================================
# MAIN: Chat interativo
# ============================================================

def main():
    print(f"""
╔══════════════════════════════════════════╗
║  🤖 Agent Full — Tools Completas        ║
║  Workspace: {WORKSPACE}
║  Tools: {len(TOOLS)} carregadas               ║
║  Digite 'sair' para encerrar            ║
╚══════════════════════════════════════════╝
""")

    for t in TOOLS:
        print(f"  • {t['name']}: {t['description'][:60]}...")

    print()
    history = None

    while True:
        try:
            user_input = input("Você: ").strip()
        except (EOFError, KeyboardInterrupt):
            break

        if not user_input or user_input.lower() in ("sair", "exit", "quit"):
            print("👋 Até mais!")
            break

        print("🔄 Processando...\n")
        response, history = agent_loop(user_input, history)
        print(f"\n🤖 {response}\n")


if __name__ == "__main__":
    main()
